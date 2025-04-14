#!/usr/bin/env python3
import os
import json
import sys
import datetime
from typing import Dict, Any, List, Optional, Union
from fastapi import FastAPI, HTTPException, Body
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt

app = FastAPI(title="LumenTrialGuide Assistant Context API")

# Advanced therapeutic area classifier with hierarchical mapping
THERAPEUTIC_AREA_KEYWORDS = {
    "Oncology": {
        "patterns": ["cancer", "oncology", "tumor", "carcinoma", "lymphoma", "leukemia", "melanoma", "sarcoma", "metastatic", "neoplasm"],
        "subtypes": {
            "Solid Tumors": ["carcinoma", "sarcoma", "melanoma", "solid tumor", "solid tumour"],
            "Hematologic": ["lymphoma", "leukemia", "myeloma", "blood cancer"],
            "Immunotherapy": ["checkpoint inhibitor", "car-t", "immunotherapy", "immune-oncology"],
            "Targeted Therapy": ["kinase inhibitor", "targeted therapy", "mutation", "EGFR", "ALK", "BRAF"]
        }
    },
    "Neurology": {
        "patterns": ["alzheimer", "dementia", "parkinsons", "epilepsy", "seizure", "neurology", "multiple sclerosis", "neurological", "brain", "cns", "central nervous system"],
        "subtypes": {
            "Neurodegenerative": ["alzheimer", "parkinson", "dementia", "lou gehrig", "huntington", "neurodegeneration"],
            "Seizure Disorders": ["epilepsy", "seizure", "convulsion"],
            "Movement Disorders": ["movement disorder", "dyskinesia", "dystonia", "ataxia"],
            "Demyelinating": ["multiple sclerosis", "ms", "demyelinating", "myelin"]
        }
    },
    "Cardiovascular": {
        "patterns": ["cardio", "heart", "vascular", "arterial", "hypertension", "cholesterol", "stroke", "thrombosis", "embolism", "coronary"],
        "subtypes": {
            "Hypertension": ["hypertension", "blood pressure", "hypertensive"],
            "Heart Failure": ["heart failure", "cardiac failure", "chf", "cardiomyopathy", "myocardial"],
            "Dyslipidemia": ["cholesterol", "lipid", "ldl", "hdl", "statin"],
            "Thrombotic": ["thrombosis", "embolism", "anti-coagulant", "anticoagulant", "clot"]
        }
    },
    "Immunology": {
        "patterns": ["inflammat", "rheumatoid", "arthritis", "autoimmune", "lupus", "psoriasis", "immune", "auto-immune"],
        "subtypes": {
            "Rheumatology": ["rheumatoid arthritis", "ra", "ankylosing spondylitis", "psoriatic arthritis", "lupus"],
            "Dermatologic": ["psoriasis", "atopic dermatitis", "eczema"],
            "GI Inflammatory": ["inflammatory bowel disease", "ibd", "crohn", "ulcerative colitis", "uc"],
            "Connective Tissue": ["lupus", "scleroderma", "sjogren", "vasculitis"]
        }
    },
    "Infectious Disease": {
        "patterns": ["infect", "hiv", "viral", "bacteria", "hepatitis", "covid", "antibiotic", "antimicrobial", "antiviral", "antibiotic"],
        "subtypes": {
            "Viral": ["hiv", "hepatitis", "herpes", "influenza", "covid", "viral"],
            "Bacterial": ["bacterial", "antibiotic", "gram-positive", "gram-negative", "mrsa"],
            "Fungal": ["fungal", "candida", "aspergillus", "antifungal"],
            "Tropical": ["malaria", "dengue", "zika", "tropical", "parasitic"]
        }
    },
    "Metabolic Disorders": {
        "patterns": ["diabetes", "insulin", "glycemic", "hyperglycemia", "obesity", "weight", "metabolic"],
        "subtypes": {
            "Diabetes": ["diabetes", "insulin", "glycemic", "hyperglycemia", "t1dm", "t2dm"],
            "Obesity": ["obesity", "weight", "overweight", "bmi", "bariatric"],
            "Lipid Disorders": ["dyslipidemia", "hyperlipidemia", "cholesterol"],
            "Rare Metabolic": ["lysosomal", "storage disorder", "fabry", "gaucher", "pompe"]
        }
    },
    "Respiratory": {
        "patterns": ["respiratory", "asthma", "copd", "lung", "pulmonary", "pneumonia", "bronchitis"],
        "subtypes": {
            "Asthma": ["asthma", "bronchial", "bronchodilator"],
            "COPD": ["copd", "chronic obstructive", "emphysema", "chronic bronchitis"],
            "Pulmonary Fibrosis": ["pulmonary fibrosis", "ipf", "fibrotic lung"],
            "Pulmonary Hypertension": ["pulmonary hypertension", "pah"]
        }
    },
    "Gastroenterology": {
        "patterns": ["gastr", "crohn", "ibs", "colitis", "bowel", "intestinal", "hepatic", "liver", "pancreatic"],
        "subtypes": {
            "IBD": ["inflammatory bowel", "crohn", "ulcerative colitis", "uc"],
            "Liver Disease": ["hepatic", "liver", "cirrhosis", "nash", "hepatitis"],
            "Functional GI": ["ibs", "functional", "constipation", "diarrhea", "motility"],
            "Pancreatic": ["pancreatic", "pancreatitis", "exocrine pancreatic insufficiency"]
        }
    },
    "Nephrology": {
        "patterns": ["renal", "kidney", "urinary", "bladder", "nephrology", "dialysis", "nephropathy"],
        "subtypes": {
            "CKD": ["chronic kidney disease", "ckd", "renal failure", "renal insufficiency"],
            "AKI": ["acute kidney injury", "aki"],
            "Glomerular": ["glomerular", "nephrotic", "glomerulonephritis"],
            "Dialysis": ["dialysis", "hemodialysis", "peritoneal"]
        }
    },
    "Rare Disease": {
        "patterns": ["rare disease", "orphan", "genetic disorder", "inherited", "congenital", "ultra-rare"],
        "subtypes": {
            "Genetic": ["genetic", "hereditary", "inherited", "chromosomal"],
            "Congenital": ["congenital", "birth defect", "developmental"],
            "Lysosomal": ["lysosomal", "storage disorder", "enzyme deficiency"],
            "Muscular": ["muscular dystrophy", "duchenne", "myopathy"]
        }
    },
    "Women's Health": {
        "patterns": ["women", "gynecology", "fertility", "contraception", "pregnancy", "menopause", "ovarian"],
        "subtypes": {
            "Reproductive": ["fertility", "infertility", "contraception", "reproduction"],
            "Pregnancy": ["pregnancy", "maternal", "obstetrics", "prenatal"],
            "Menopause": ["menopause", "postmenopausal", "hot flashes"],
            "Gynecologic": ["gynecological", "uterine", "ovarian", "cervical"]
        }
    },
    "Psychiatry": {
        "patterns": ["depression", "anxiety", "bipolar", "schizophrenia", "psychiatric", "mental health", "ptsd", "adhd"],
        "subtypes": {
            "Mood Disorders": ["depression", "bipolar", "major depressive", "mdd"],
            "Anxiety Disorders": ["anxiety", "gad", "panic", "social anxiety", "ptsd"],
            "Psychotic Disorders": ["schizophrenia", "psychosis", "psychotic"],
            "Neurodevelopmental": ["adhd", "attention deficit", "autism", "asd"]
        }
    }
}

def classify_therapeutic_area(indication_text: str) -> Dict[str, Any]:
    """
    Advanced therapeutic area classifier that provides both main area and more specific subtype
    
    Args:
        indication_text: Text describing the indication
        
    Returns:
        Dict with main therapeutic area, subtype, and confidence scores
    """
    if not indication_text or not isinstance(indication_text, str):
        return {
            "main_area": "Unknown",
            "subtype": "Unknown",
            "main_confidence": 0.0,
            "subtype_confidence": 0.0
        }
    
    indication_lower = indication_text.lower()
    
    # Score for each therapeutic area based on keyword matches
    area_scores = {}
    subtype_scores = {}
    
    # Calculate scores for each therapeutic area and its subtypes
    for area, data in THERAPEUTIC_AREA_KEYWORDS.items():
        # Calculate main area score
        patterns = data["patterns"]
        area_score = sum(1 for pattern in patterns if pattern in indication_lower)
        
        if area_score > 0:
            area_scores[area] = area_score
            
            # Calculate subtype scores if main area has a match
            subtype_scores[area] = {}
            for subtype, subpatterns in data["subtypes"].items():
                subtype_score = sum(1 for pattern in subpatterns if pattern in indication_lower)
                if subtype_score > 0:
                    subtype_scores[area][subtype] = subtype_score
    
    # Determine main therapeutic area
    if not area_scores:
        return {
            "main_area": "Other",
            "subtype": "General",
            "main_confidence": 0.5,
            "subtype_confidence": 0.0
        }
    
    # Sort areas by score
    sorted_areas = sorted(area_scores.items(), key=lambda x: x[1], reverse=True)
    main_area = sorted_areas[0][0]
    main_score = sorted_areas[0][1]
    
    # Calculate confidence (normalize by max possible score)
    main_confidence = min(main_score / len(THERAPEUTIC_AREA_KEYWORDS[main_area]["patterns"]), 1.0)
    
    # Determine subtype for main area
    subtype = "General"
    subtype_confidence = 0.0
    
    if main_area in subtype_scores and subtype_scores[main_area]:
        sorted_subtypes = sorted(subtype_scores[main_area].items(), key=lambda x: x[1], reverse=True)
        subtype = sorted_subtypes[0][0]
        subtype_score = sorted_subtypes[0][1]
        
        # Calculate subtype confidence
        max_possible_subtype_score = len(THERAPEUTIC_AREA_KEYWORDS[main_area]["subtypes"][subtype])
        subtype_confidence = min(subtype_score / max_possible_subtype_score, 1.0)
    
    return {
        "main_area": main_area,
        "subtype": subtype,
        "main_confidence": main_confidence,
        "subtype_confidence": subtype_confidence
    }

def get_regulatory_considerations(therapeutic_area: str, subtype: str) -> List[str]:
    """
    Get specific regulatory considerations based on therapeutic area and subtype
    
    Args:
        therapeutic_area: Main therapeutic area
        subtype: Specific subtype within the therapeutic area
        
    Returns:
        List of relevant regulatory considerations
    """
    # Sophisticated regulatory insights based on therapeutic area
    considerations = {
        "Oncology": {
            "General": [
                "Consider accelerated approval pathways based on surrogate endpoints like ORR or PFS",
                "Ensure collection of patient-reported outcomes for symptomatic benefit",
                "Plan for post-marketing commitments to verify clinical benefit"
            ],
            "Solid Tumors": [
                "RECIST criteria should be used for response evaluation unless justified otherwise",
                "Consider stratification by biomarker status and prior lines of therapy",
                "Include QoL assessments to support labeling claims on symptom improvement"
            ],
            "Hematologic": [
                "Consider MRD (minimal residual disease) as an endpoint for hematologic malignancies",
                "Plan for long-term safety monitoring particularly for CAR-T and other cell therapies",
                "FDA recommends including transplant-eligible and ineligible populations where appropriate"
            ],
            "Immunotherapy": [
                "Account for potential pseudoprogression in response criteria",
                "Include immune-related adverse event management guidelines",
                "Consider novel endpoints such as TMDD (treatment beyond progression)"
            ]
        },
        "Neurology": {
            "General": [
                "Include validated cognitive assessment tools relevant to the specific condition",
                "Consider caregiver burden assessments and quality of life measures",
                "FDA expects functional outcomes in addition to symptomatic improvement"
            ],
            "Neurodegenerative": [
                "Include biomarkers (e.g., CSF markers, PET imaging) to demonstrate target engagement",
                "Plan for extended follow-up periods to demonstrate disease modification",
                "Consider adaptive trial designs with interim analyses for early signals"
            ],
            "Seizure Disorders": [
                "Time to first seizure and seizure frequency are acceptable primary endpoints",
                "Include standardized seizure diaries and validated classification systems",
                "Consider video EEG monitoring for objective seizure assessment"
            ]
        },
        "Cardiovascular": {
            "General": [
                "FDA expects hard outcome endpoints (MACE, CV death) for pivotal CV trials",
                "Design with appropriate safety monitoring including independent DSMB",
                "Consider risk-based approaches for monitoring and data collection"
            ],
            "Heart Failure": [
                "Time to CV death or HF hospitalization is the gold standard endpoint",
                "Include KCCQ and other validated PRO measures to support symptomatic benefit claims",
                "FDA guidance supports NT-proBNP as a surrogate for early phase studies only"
            ],
            "Hypertension": [
                "Change in systolic/diastolic blood pressure is acceptable for registration",
                "Include 24-hour ambulatory blood pressure monitoring",
                "Recent FDA guidance suggests focus on sustained BP control rather than single timepoints"
            ]
        }
    }
    
    # Return considerations for the specific area and subtype
    default_considerations = [
        "Follow ICH E8(R1) guidelines for general considerations of clinical studies",
        "Ensure appropriate diversity in study populations",
        "Include clear statistical analysis plan with pre-specified endpoints"
    ]
    
    area_considerations = considerations.get(therapeutic_area, {}).get("General", [])
    specific_considerations = considerations.get(therapeutic_area, {}).get(subtype, [])
    
    return specific_considerations + area_considerations + default_considerations

def get_trial_design_insights(therapeutic_area: str, subtype: str) -> Dict[str, Any]:
    """
    Get specific trial design insights based on therapeutic area and subtype
    
    Args:
        therapeutic_area: Main therapeutic area
        subtype: Specific subtype within the therapeutic area
        
    Returns:
        Dict with trial design recommendations
    """
    # Sophisticated trial design insights
    design_insights = {
        "Oncology": {
            "General": {
                "recommended_designs": ["Randomized controlled", "Single-arm for rare tumors"],
                "control_considerations": "Consider synthetic control arms for rare tumors",
                "endpoint_strategy": "Progression-free survival (PFS) for early lines, Overall survival (OS) for later lines",
                "sample_size_factors": "Effect size typically 0.65-0.75 HR for PFS, 0.7-0.8 HR for OS",
                "biomarker_strategy": "Consider enrichment design for targeted therapies"
            },
            "Immunotherapy": {
                "recommended_designs": ["Randomized with crossover consideration", "Simon 2-stage for early development"],
                "control_considerations": "SoC with careful planning for crossover",
                "endpoint_strategy": "Consider 6-month landmark analyses to account for delayed effect",
                "sample_size_factors": "Plan for delayed separation of curves, potentially larger sample sizes",
                "biomarker_strategy": "Include TMB, PD-L1, MSI status as stratification factors"
            }
        },
        "Neurology": {
            "General": {
                "recommended_designs": ["Double-blind, placebo-controlled", "Delayed-start design"],
                "control_considerations": "Placebo control essential for symptomatic endpoints",
                "endpoint_strategy": "Combination of symptom scores and functional outcomes",
                "sample_size_factors": "High variability requires larger samples, consider 20% increase",
                "biomarker_strategy": "Include imaging and fluid biomarkers as secondary endpoints"
            },
            "Neurodegenerative": {
                "recommended_designs": ["Delayed-start", "Long-term extension with interim analysis"],
                "control_considerations": "Placebo control with ethical consideration for duration",
                "endpoint_strategy": "Rate of decline in functional/cognitive measures",
                "sample_size_factors": "Account for variable progression rates, consider 30% larger sample",
                "biomarker_strategy": "Include target engagement and disease progression biomarkers"
            }
        }
    }
    
    # Default design insights
    default_insights = {
        "recommended_designs": ["Randomized controlled trial"],
        "control_considerations": "Standard of care comparison",
        "endpoint_strategy": "Clinically meaningful endpoints with regulatory precedent",
        "sample_size_factors": "Based on expected effect size with 80-90% power",
        "biomarker_strategy": "Exploratory biomarkers for mechanism validation"
    }
    
    # Get specific insights or fall back to defaults
    area_insights = design_insights.get(therapeutic_area, {}).get("General", default_insights)
    specific_insights = design_insights.get(therapeutic_area, {}).get(subtype, area_insights)
    
    return specific_insights if specific_insights != area_insights else area_insights

@app.get("/api/context/assistant-csr/{csr_id}")
def get_assistant_context_from_csr(csr_id: str):
    """
    Get contextual information from a CSR for the assistant
    
    Args:
        csr_id: The CSR ID to get context from
        
    Returns:
        Dict with CSR context information and assistant memory prompt
    """
    csr_path = f"/mnt/data/lumen_reports_backend/intelligence_db/{csr_id}.json"
    if not os.path.exists(csr_path):
        csr_path = f"./data/intelligence_db/{csr_id}.json"  # Fallback path
        if not os.path.exists(csr_path):
            raise HTTPException(status_code=404, detail=f"CSR not found: {csr_id}")

    with open(csr_path, "r") as f:
        csr_data = json.load(f)

    # Extract key information
    molecule = csr_data.get("meta", {}).get("molecule", "unknown molecule")
    moa = csr_data.get("pharmacology", {}).get("moa_explained") or csr_data.get("meta", {}).get("moa", "unknown MoA")
    indication = csr_data.get("indication", "unspecified indication")
    endpoint = csr_data.get("efficacy", {}).get("primary", ["unspecified endpoint"])[0]
    rationale = csr_data.get("semantic", {}).get("design_rationale", "no rationale provided")
    model = csr_data.get("stats_traceability", {}).get("primary_model", "unspecified model")
    
    # Enhanced therapeutic area classification
    ta_classification = classify_therapeutic_area(indication)
    main_area = ta_classification["main_area"]
    subtype = ta_classification["subtype"]
    
    # Get regulatory considerations based on therapeutic area
    regulatory_considerations = get_regulatory_considerations(main_area, subtype)
    
    # Get trial design insights
    design_insights = get_trial_design_insights(main_area, subtype)

    # Create enhanced memory prompt
    memory_prompt = (
        f"This planning session was launched based on CSR `{csr_id}`.\n"
        f"Molecule: {molecule}\n"
        f"Mechanism of Action: {moa}\n"
        f"Indication: {indication}\n"
        f"Therapeutic Area: {main_area} ({subtype})\n"
        f"Primary Endpoint: {endpoint}\n"
        f"Design Rationale: {rationale}\n"
        f"Statistical Model: {model}\n\n"
        f"Key Regulatory Considerations:\n"
        f"- {regulatory_considerations[0]}\n"
        f"- {regulatory_considerations[1]}\n"
        f"Use this contextual evidence to inform all assistant responses, IND drafting, and protocol validations."
    )

    return {
        "csr_id": csr_id,
        "molecule": molecule,
        "moa": moa,
        "indication": indication,
        "therapeutic_area": {
            "main_area": main_area,
            "subtype": subtype,
            "confidence": {
                "main": ta_classification["main_confidence"],
                "subtype": ta_classification["subtype_confidence"]
            }
        },
        "endpoint": endpoint,
        "rationale": rationale,
        "model": model,
        "regulatory_considerations": regulatory_considerations,
        "design_insights": design_insights,
        "assistant_memory": memory_prompt
    }

@app.post("/api/planner/generate-ind")
async def generate_ind_summary(request: Dict[str, Any]):
    """
    Generate an IND summary based on protocol and CSR context
    
    Args:
        request: Dict containing protocol and CSR context
        
    Returns:
        Dict with generated IND summary
    """
    protocol = request.get("protocol", "")
    session_id = request.get("sessionId", "")
    csr_context = request.get("csrContext", {})
    
    if not protocol:
        raise HTTPException(status_code=400, detail="Protocol is required")
    
    # Extract key information
    indication = csr_context.get("indication", "")
    phase = csr_context.get("phase", "")
    molecule = csr_context.get("drugName", "")
    
    # Enhanced therapeutic area classification if indication is available
    ta_classification = classify_therapeutic_area(indication)
    main_area = ta_classification["main_area"]
    subtype = ta_classification["subtype"]
    
    # Get regulatory considerations
    regulatory_considerations = get_regulatory_considerations(main_area, subtype)
    
    # Use advanced fields from csr_context if available
    design_rationale = csr_context.get("design_rationale", "")
    moa = csr_context.get("moa", "")
    primary_model = csr_context.get("primary_model", "")
    primary_endpoint = csr_context.get("primary_endpoint", "")
    
    # Get trial design insights
    design_insights = get_trial_design_insights(main_area, subtype)
    recommended_design = design_insights.get("recommended_designs", [""])[0] if design_insights.get("recommended_designs") else ""
    endpoint_strategy = design_insights.get("endpoint_strategy", "")
    
    # Generate IND summary template based on therapeutic area and available context
    ind_summary = f"""# INVESTIGATIONAL NEW DRUG APPLICATION SUMMARY

## 1. INTRODUCTION

This Investigational New Drug (IND) application pertains to the clinical investigation of {molecule} for the treatment of {indication}.

**Therapeutic Area**: {main_area} ({subtype})

## 2. DRUG SUBSTANCE AND FORMULATION

{molecule} is a {"novel " if moa else ""}compound that {"functions through " + moa if moa else "is being investigated for therapeutic potential"}.

## 3. PRECLINICAL PHARMACOLOGY AND TOXICOLOGY

### Mechanism of Action
{moa if moa else "The detailed mechanism of action will be described in the complete IND submission."}

{"### Therapeutic Area-Specific Considerations" if regulatory_considerations else ""}
{("- " + regulatory_considerations[0]) if regulatory_considerations else ""}
{("- " + regulatory_considerations[1]) if regulatory_considerations else ""}

## 4. CLINICAL PROTOCOL SUMMARY

### Study Design
{"Based on precedent in " + main_area + " (" + subtype + "), a " + recommended_design + " design is proposed." if recommended_design else "The study design was developed based on standard approaches for this indication and phase of development."}

### Study Design Rationale
{design_rationale if design_rationale else "The study design was developed based on standard approaches for this indication and phase of development."}

### Primary Endpoint
{primary_endpoint if primary_endpoint else "The primary endpoint will be specified in the full protocol."}
{("Endpoint Strategy: " + endpoint_strategy) if endpoint_strategy else ""}

### Statistical Approach
{primary_model if primary_model else "Standard statistical methodologies appropriate for this therapeutic area will be employed."}

## 5. INVESTIGATOR AND FACILITIES INFORMATION

[To be completed with investigator and facility details]

## 6. CHEMISTRY, MANUFACTURING, AND CONTROLS

[Detailed CMC information will be provided in the complete submission]

"""

    return {
        "success": True,
        "content": ind_summary,
        "sessionId": session_id
    }

@app.post("/api/export/ind-docx-with-context")
def export_ind_with_context(data: Dict = Body(...)):
    """
    Export IND summary with CSR context as a DOCX document
    
    Args:
        data: Dict containing session_id and csr_id
        
    Returns:
        DOCX file download response
    """
    session_id = data.get("session_id", "default_session")
    csr_id = data.get("csr_id", session_id)

    archive_dir = f"/mnt/data/lumen_reports_backend/sessions/{session_id}"
    os.makedirs(archive_dir, exist_ok=True)

    # Load CSR data
    csr_path = f"/mnt/data/lumen_reports_backend/intelligence_db/{csr_id}.json"
    if not os.path.exists(csr_path):
        csr_path = f"./data/intelligence_db/{csr_id}.json"  # Fallback path
        if not os.path.exists(csr_path):
            raise HTTPException(status_code=404, detail=f"CSR file not found: {csr_id}")

    with open(csr_path, "r") as f:
        csr_data = json.load(f)

    # Extract key information
    molecule = csr_data.get("meta", {}).get("molecule", "unspecified")
    moa = csr_data.get("pharmacology", {}).get("moa_explained") or csr_data.get("meta", {}).get("moa", "unspecified")
    indication = csr_data.get("indication", "unspecified indication")
    endpoint = csr_data.get("efficacy", {}).get("primary", [""])[0]
    rationale = csr_data.get("semantic", {}).get("design_rationale", "")
    model = csr_data.get("stats_traceability", {}).get("primary_model", "")

    # Enhanced therapeutic area classification
    ta_classification = classify_therapeutic_area(indication)
    main_area = ta_classification["main_area"]
    subtype = ta_classification["subtype"]

    # Get regulatory considerations
    regulatory_considerations = get_regulatory_considerations(main_area, subtype)[:2]  # Get top 2 considerations
    
    # Get trial design insights
    design_insights = get_trial_design_insights(main_area, subtype)

    # Create DOCX
    doc = Document()
    doc.add_heading("LumenTrialGuide.AI – IND Module 2.5", 0)
    doc.add_paragraph(f"Study ID: {session_id}")
    doc.add_paragraph(f"CSR Source: {csr_id}")
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("-" * 80)

    doc.add_heading("📄 IND Summary Context", level=1)
    doc.add_paragraph(f"Molecule: {molecule}")
    doc.add_paragraph(f"Mechanism of Action: {moa}")
    doc.add_paragraph(f"Indication: {indication}")
    doc.add_paragraph(f"Therapeutic Area: {main_area} ({subtype})")
    doc.add_paragraph(f"Primary Endpoint: {endpoint}")
    doc.add_paragraph(f"Design Rationale: {rationale}")
    doc.add_paragraph(f"Statistical Model: {model}")
    
    if regulatory_considerations:
        doc.add_heading("🔍 Regulatory Considerations", level=2)
        for consideration in regulatory_considerations:
            doc.add_paragraph(f"• {consideration}")
    
    if design_insights:
        doc.add_heading("🧩 Trial Design Strategy", level=2)
        for key, value in design_insights.items():
            if isinstance(value, list):
                value_text = ", ".join(value)
            else:
                value_text = value
            doc.add_paragraph(f"• {key.replace('_', ' ').title()}: {value_text}")

    doc.add_heading("🧠 AI-Generated Narrative", level=1)
    doc.add_paragraph(
        f"This protocol was derived using CSR precedent from {csr_id}, involving a {moa}. "
        f"The study focuses on {endpoint}, with rationale: {rationale}. "
        f"The statistical model utilized is {model}, aligning with similar regulatory-approved designs in {main_area} ({subtype})."
    )

    doc_path = os.path.join(archive_dir, "ind_module_with_context.docx")
    doc.save(doc_path)

    return FileResponse(doc_path, filename="ind_module_with_context.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/api/planning/init")
def initialize_planning_session(data: Dict = Body(...)):
    """
    Initialize a planning session with auto-generated IND document from CSR
    
    Args:
        data: Dict containing session_id and csr_id
        
    Returns:
        Dict with initialization status
    """
    session_id = data.get("session_id", "default_session")
    csr_id = data.get("csr_id", None)
    archive_dir = f"/mnt/data/lumen_reports_backend/sessions/{session_id}"
    os.makedirs(archive_dir, exist_ok=True)

    ind_path = os.path.join(archive_dir, "ind_module_with_context.docx")
    if csr_id and not os.path.exists(ind_path):
        csr_path = f"/mnt/data/lumen_reports_backend/intelligence_db/{csr_id}.json"
        if not os.path.exists(csr_path):
            csr_path = f"./data/intelligence_db/{csr_id}.json"  # Fallback path
            if not os.path.exists(csr_path):
                return {"status": "error", "message": f"CSR file not found: {csr_id}"}

        with open(csr_path, "r") as f:
            csr_data = json.load(f)

        # Extract key information
        molecule = csr_data.get("meta", {}).get("molecule", "unspecified")
        moa = csr_data.get("pharmacology", {}).get("moa_explained") or csr_data.get("meta", {}).get("moa", "unspecified")
        indication = csr_data.get("indication", "unspecified indication")
        endpoint = csr_data.get("efficacy", {}).get("primary", [""])[0]
        rationale = csr_data.get("semantic", {}).get("design_rationale", "")
        model = csr_data.get("stats_traceability", {}).get("primary_model", "")

        # Enhanced therapeutic area classification
        ta_classification = classify_therapeutic_area(indication)
        main_area = ta_classification["main_area"]
        subtype = ta_classification["subtype"]

        # Get regulatory considerations
        regulatory_considerations = get_regulatory_considerations(main_area, subtype)[:2]  # Get top 2 considerations
        
        # Get trial design insights
        design_insights = get_trial_design_insights(main_area, subtype)

        # Create DOCX
        doc = Document()
        doc.add_heading("LumenTrialGuide.AI – IND Module 2.5", 0)
        doc.add_paragraph(f"Study ID: {session_id}")
        doc.add_paragraph(f"CSR Source: {csr_id}")
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("-" * 80)

        doc.add_heading("📄 IND Summary Context", level=1)
        doc.add_paragraph(f"Molecule: {molecule}")
        doc.add_paragraph(f"Mechanism of Action: {moa}")
        doc.add_paragraph(f"Indication: {indication}")
        doc.add_paragraph(f"Therapeutic Area: {main_area} ({subtype})")
        doc.add_paragraph(f"Primary Endpoint: {endpoint}")
        doc.add_paragraph(f"Design Rationale: {rationale}")
        doc.add_paragraph(f"Statistical Model: {model}")
        
        if regulatory_considerations:
            doc.add_heading("🔍 Regulatory Considerations", level=2)
            for consideration in regulatory_considerations:
                doc.add_paragraph(f"• {consideration}")
        
        if design_insights:
            doc.add_heading("🧩 Trial Design Strategy", level=2)
            for key, value in design_insights.items():
                if isinstance(value, list):
                    value_text = ", ".join(value)
                else:
                    value_text = value
                doc.add_paragraph(f"• {key.replace('_', ' ').title()}: {value_text}")

        doc.add_heading("🧠 AI-Generated Narrative", level=1)
        doc.add_paragraph(
            f"This protocol was derived using CSR precedent from {csr_id}, involving a {moa}. "
            f"The study focuses on {endpoint}, with rationale: {rationale}. "
            f"The statistical model utilized is {model}, aligning with similar regulatory-approved designs in {main_area} ({subtype})."
        )

        doc.save(ind_path)

    return {"status": "initialized", "session_id": session_id, "csr_id": csr_id}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)