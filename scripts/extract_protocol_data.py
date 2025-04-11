#!/usr/bin/env python3
"""
Protocol Data Extraction Script
This script extracts structured data from clinical trial protocol documents (.docx or .txt)
"""

import sys
import os
import re
import json
import argparse
from typing import Dict, Any, List, Optional

# For docx file parsing
try:
    import docx
except ImportError:
    print("Warning: python-docx not installed. DOCX parsing will not be available.")

def clean_text(text: str) -> str:
    """Clean and normalize text for better extraction"""
    text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII characters
    return text.strip()

def extract_from_docx(file_path: str) -> str:
    """Extract text content from a .docx file"""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def extract_from_txt(file_path: str) -> str:
    """Extract text content from a .txt file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        print(f"Error extracting text from TXT: {e}")
        return ""

def extract_text_from_file(file_path: str) -> str:
    """Extract text from either docx or txt file"""
    _, ext = os.path.splitext(file_path)
    if ext.lower() == '.docx':
        return extract_from_docx(file_path)
    elif ext.lower() == '.txt':
        return extract_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def find_in_text(text: str, patterns: List[str], is_integer: bool = False) -> Optional[str]:
    """Find a pattern in text and extract the matched group"""
    for pattern in patterns:
        if is_integer:
            match = re.search(rf"{pattern}\s*[:\-]?\s*(\d+)", text, re.IGNORECASE | re.MULTILINE)
        else:
            match = re.search(rf"{pattern}\s*[:\-]?\s*(.*?)(?:\.|\n|$)", text, re.IGNORECASE | re.MULTILINE)
        
        if match:
            return match.group(1).strip()
    return None

def find_integer(text: str, patterns: List[str]) -> Optional[int]:
    """Find an integer value in text based on patterns"""
    result = find_in_text(text, patterns, True)
    if result:
        try:
            return int(result)
        except ValueError:
            pass
    return None

def find_float(text: str, patterns: List[str]) -> Optional[float]:
    """Find a float value in text based on patterns"""
    for pattern in patterns:
        match = re.search(rf"{pattern}\s*[:\-]?\s*(\d+\.?\d*|\.\d+)%?", text, re.IGNORECASE | re.MULTILINE)
        if match:
            value = match.group(1).strip()
            try:
                val = float(value)
                # If the pattern contains 'percent' or '%' and the value is ≤ 1, 
                # it's likely already in decimal form
                if 'percent' in pattern.lower() or '%' in pattern:
                    if val > 1:
                        val = val / 100
                return val
            except ValueError:
                pass
    return None

def extract_list(text: str, pattern: str) -> List[str]:
    """Extract a list of items from text"""
    section = re.search(rf"{pattern}(.*?)(?:\n\n|\n[A-Z])", text, re.IGNORECASE | re.DOTALL)
    if not section:
        return []
    
    items = re.findall(r'(?:^|\n)[\•\-\*\s]*(\d+\.\s*|[a-z]\)\s*|\([a-z]\)\s*)?([^•\n][^\n]+)', section.group(1))
    return [item[1].strip() for item in items if item[1].strip()]

def find_phase(text: str) -> str:
    """Extract study phase"""
    phase_pattern = r"Phase\s*((?:1|2|3)(?:\/(?:1|2|3))?[a-zA-Z]*)"
    phase_match = re.search(phase_pattern, text, re.IGNORECASE)
    
    if phase_match:
        phase = phase_match.group(1).strip()
        # Normalize phase format
        if '/' in phase:
            parts = phase.split('/')
            return f"Phase {parts[0]}/{parts[1]}"
        return f"Phase {phase}"
    
    # Try alternative patterns
    if re.search(r"Phase\s*I[IVb]*", text, re.IGNORECASE):
        match = re.search(r"Phase\s*(I[IVb]*)", text, re.IGNORECASE)
        if match:
            phase_roman = match.group(1)
            # Convert Roman numerals to Arabic
            if phase_roman == "I":
                return "Phase 1"
            elif phase_roman in ["II", "IIa", "IIb"]:
                suffix = ""
                if "a" in phase_roman:
                    suffix = "a"
                elif "b" in phase_roman:
                    suffix = "b"
                return f"Phase 2{suffix}"
            elif phase_roman in ["III", "IIIa", "IIIb"]:
                suffix = ""
                if "a" in phase_roman:
                    suffix = "a"
                elif "b" in phase_roman:
                    suffix = "b"
                return f"Phase 3{suffix}"
            elif phase_roman == "IV":
                return "Phase 4"
    
    return "Unknown"

def extract_endpoints(text: str) -> Dict[str, List[str]]:
    """Extract primary and secondary endpoints"""
    # Try to find a primary endpoint section
    primary_patterns = [
        r"primary endpoint[s]?.*?:?\s*(.*?)(?:\.|$)",
        r"primary outcome[s]?.*?:?\s*(.*?)(?:\.|$)",
        r"primary objective[s]?.*?:?\s*(.*?)(?:\.|$)"
    ]
    
    primary_endpoints = []
    for pattern in primary_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            primary_endpoints.append(match.group(1).strip())
    
    # Try to find secondary endpoint section
    secondary_patterns = [
        r"secondary endpoint[s]?.*?:?\s*(.*?)(?:\.|$)",
        r"secondary outcome[s]?.*?:?\s*(.*?)(?:\.|$)",
        r"secondary objective[s]?.*?:?\s*(.*?)(?:\.|$)"
    ]
    
    secondary_endpoints = []
    for pattern in secondary_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            secondary_endpoints.append(match.group(1).strip())
    
    return {
        "primary": primary_endpoints,
        "secondary": secondary_endpoints
    }

def extract_arms(text: str) -> List[str]:
    """Extract study arms"""
    arms_patterns = [
        r"(?:study|treatment) arms?:?\s*(.*?)(?:\n\n|\n[A-Z])",
        r"(?:study|treatment) groups?:?\s*(.*?)(?:\n\n|\n[A-Z])"
    ]
    
    for pattern in arms_patterns:
        section = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if section:
            items = re.findall(r'(?:^|\n)[\•\-\*\s]*(?:\d+\.\s*)?([^•\n][^\n]+)', section.group(1))
            if items:
                return [item.strip() for item in items if item.strip()]
    
    # If no explicit arms section, try to infer from text
    arms = []
    
    # Look for placebo mentions
    if re.search(r"placebo.{0,30}(?:control|arm|group)", text, re.IGNORECASE):
        arms.append("Placebo")
    
    # Look for active control mentions
    if re.search(r"active.{0,30}(?:control|arm|group)", text, re.IGNORECASE):
        arms.append("Active Control")
    
    # Look for intervention/treatment arm
    if re.search(r"(?:intervention|treatment|experimental).{0,30}(?:arm|group)", text, re.IGNORECASE):
        arms.append("Treatment")
    
    return arms

def extract_blinding(text: str) -> str:
    """Extract study blinding information"""
    if re.search(r"double.{0,15}blind", text, re.IGNORECASE):
        return "Double-blind"
    elif re.search(r"single.{0,15}blind", text, re.IGNORECASE):
        return "Single-blind"
    elif re.search(r"triple.{0,15}blind", text, re.IGNORECASE):
        return "Triple-blind"
    elif re.search(r"open.{0,15}label", text, re.IGNORECASE):
        return "Open-label"
    elif re.search(r"not.{0,15}blinded", text, re.IGNORECASE):
        return "Not blinded"
    
    return "Unknown"

def extract_protocol_data(text: str) -> Dict[str, Any]:
    """Extract all protocol data from text"""
    # Clean the text
    text = clean_text(text)
    
    # Extract title
    title = find_in_text(text, ["study title", "protocol title", "title"])
    if not title:
        title = "Untitled Protocol"
    
    # Extract phase
    phase = find_phase(text)
    
    # Extract indication/disease
    indication = find_in_text(text, [
        "indication", "disease", "condition", "population", "study population"
    ])
    
    # Extract sample size
    sample_size = find_integer(text, [
        "sample size", "n=", "subjects", "patients", "participants", "total enrollment"
    ])
    
    # Extract duration
    duration_weeks = find_integer(text, [
        "duration.{0,20}weeks", "weeks.{0,20}duration", "study duration", "treatment duration", "follow.{0,20}up"
    ])
    
    # Extract dropout rate/attrition
    dropout_rate = find_float(text, [
        "dropout rate", "attrition rate", "discontinuation rate", "withdrawal rate", "anticipated dropout percent"
    ])
    if dropout_rate is None:
        dropout_rate = 0.15  # Default assumption
    
    # Extract endpoints
    endpoints = extract_endpoints(text)
    
    # Extract study design
    study_design = find_in_text(text, ["study design", "design"])
    
    # Extract arms
    arms = extract_arms(text)
    
    # Extract blinding
    blinding = extract_blinding(text)
    
    # Extract inclusion criteria
    inclusion_criteria = extract_list(text, "inclusion criteria")
    
    # Extract exclusion criteria
    exclusion_criteria = extract_list(text, "exclusion criteria")
    
    # Extract statistical plan
    statistical_plan = find_in_text(text, ["statistical analysis", "statistical methods", "statistics"])
    
    # Extract comparator information
    comparator = None
    if re.search(r"compared to|versus|vs\.?|comparator", text, re.IGNORECASE):
        comparator = find_in_text(text, ["compared to", "versus", "vs\.?", "comparator"])
    
    # Extract control type
    control_type = None
    if re.search(r"placebo.{0,30}(?:control|arm|group)", text, re.IGNORECASE):
        control_type = "Placebo"
    elif re.search(r"active.{0,30}(?:control|arm|group)", text, re.IGNORECASE):
        control_type = "Active"
    elif re.search(r"no.{0,10}(?:control|comparator)", text, re.IGNORECASE):
        control_type = "None"
    
    return {
        "title": title,
        "phase": phase,
        "indication": indication,
        "sample_size": sample_size or 0,
        "duration_weeks": duration_weeks or 0,
        "dropout_rate": dropout_rate,
        "primary_endpoints": endpoints["primary"],
        "secondary_endpoints": endpoints["secondary"],
        "study_design": study_design,
        "arms": arms,
        "blinding": blinding,
        "inclusion_criteria": inclusion_criteria,
        "exclusion_criteria": exclusion_criteria,
        "statistical_plan": statistical_plan,
        "comparator": comparator,
        "control_type": control_type
    }

def main():
    parser = argparse.ArgumentParser(description="Extract structured data from clinical trial protocol documents")
    parser.add_argument("input_file", help="Path to protocol document (.docx or .txt)")
    parser.add_argument("output_file", help="Path to output JSON file")
    
    args = parser.parse_args()
    
    # Extract text from file
    text = extract_text_from_file(args.input_file)
    
    if not text:
        print("Error: Failed to extract text from file")
        sys.exit(1)
    
    # Extract protocol data
    protocol_data = extract_protocol_data(text)
    
    # Write output to JSON file
    with open(args.output_file, 'w', encoding='utf-8') as f:
        json.dump(protocol_data, f, indent=2)
    
    print(f"Protocol data extracted successfully to {args.output_file}")

if __name__ == "__main__":
    main()