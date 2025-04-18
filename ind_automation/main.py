#!/usr/bin/env python3
"""
IND Automation FastAPI Service

This module provides a FastAPI service for generating FDA Investigational New Drug (IND)
application forms and related documents.
"""

import os
import sys
import json
import uuid
from datetime import datetime
from typing import Dict, List, Optional, Any, Union
from io import BytesIO

# FastAPI imports
from fastapi import FastAPI, HTTPException, UploadFile, File, Body, Query
from fastapi.responses import StreamingResponse, FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field

# Local imports
from .templates import (
    generate_form_1571,
    generate_form_1572,
    generate_form_3674,
    generate_cover_letter
)

# Define data models
class SpecificationData(BaseModel):
    parameter: str
    limit: str
    result: str

class StabilityData(BaseModel):
    timepoint: str
    result: str

class Module3Data(BaseModel):
    drug_name: str
    manufacturing_site: str
    batch_number: str
    specifications: List[SpecificationData]
    stability_data: List[StabilityData]

class ProjectInfo(BaseModel):
    id: str
    name: str

class ProjectMetadata(BaseModel):
    sponsor_name: Optional[str] = None
    sponsor_address: Optional[str] = None
    sponsor_phone: Optional[str] = None
    ind_number: Optional[str] = None
    drug_name: Optional[str] = None
    indication: Optional[str] = None
    protocol_number: Optional[str] = None
    protocol_title: Optional[str] = None
    phase: Optional[str] = None
    submission_date: Optional[str] = None
    nct_number: Optional[str] = None
    principal_investigator_name: Optional[str] = None
    investigator_address: Optional[str] = None
    investigator_phone: Optional[str] = None
    irb_name: Optional[str] = None
    irb_address: Optional[str] = None
    clinical_lab_name: Optional[str] = None
    clinical_lab_address: Optional[str] = None
    research_facility_name: Optional[str] = None
    research_facility_address: Optional[str] = None
    subinvestigators: Optional[str] = None
    contact_name: Optional[str] = None
    contact_email: Optional[str] = None
    contact_phone: Optional[str] = None
    authorizer_name: Optional[str] = None
    authorizer_title: Optional[str] = None
    certifier_name: Optional[str] = None
    certifier_title: Optional[str] = None
    certifier_address: Optional[str] = None
    certifier_email: Optional[str] = None
    certifier_phone: Optional[str] = None
    certifier_fax: Optional[str] = None
    serial_number: Optional[str] = None

# Create FastAPI app
app = FastAPI(
    title="IND Automation API",
    description="API for generating FDA IND application forms and related documents",
    version="2.0.0"
)

# Configure CORS to allow requests from the main application
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, this should be restricted
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Root endpoint
@app.get("/")
def read_root():
    """Root endpoint providing service information"""
    return {
        "service": "IND Automation API",
        "version": "2.0.0",
        "status": "running",
        "endpoints": {
            "/health": "Health check endpoint",
            "/projects": "List available projects",
            "/generate/form1571": "Generate FDA Form 1571",
            "/generate/form1572": "Generate FDA Form 1572",
            "/generate/form3674": "Generate FDA Form 3674",
            "/generate/cover-letter": "Generate Cover Letter",
            "/generate/module3": "Generate Module 3 document"
        }
    }

# Health check endpoint
@app.get("/health")
def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

# List projects endpoint (stub for now)
@app.get("/projects")
def list_projects():
    """List available projects from Benchling (stub)"""
    # In the future, this will connect to Benchling API
    sample_projects = [
        {"id": "PRJ001", "name": "Oncology Drug X Phase I"},
        {"id": "PRJ002", "name": "Cardiovascular Drug Y Phase II"},
        {"id": "PRJ003", "name": "Immunotherapy Z Phase I/II"}
    ]
    return {"projects": sample_projects}

# Generate Form 1571
@app.post("/generate/form1571")
def generate_1571(data: ProjectMetadata):
    """Generate FDA Form 1571 (Investigational New Drug Application)"""
    try:
        # Generate the document using the template
        document_buffer = generate_form_1571(data.dict(exclude_none=True))
        
        # Return the document as a file download
        return StreamingResponse(
            document_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Form1571_{datetime.now().strftime('%Y%m%d')}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating Form 1571: {str(e)}")

# Generate Form 1572
@app.post("/generate/form1572")
def generate_1572(data: ProjectMetadata):
    """Generate FDA Form 1572 (Statement of Investigator)"""
    try:
        # Generate the document using the template
        document_buffer = generate_form_1572(data.dict(exclude_none=True))
        
        # Return the document as a file download
        return StreamingResponse(
            document_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Form1572_{datetime.now().strftime('%Y%m%d')}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating Form 1572: {str(e)}")

# Generate Form 3674
@app.post("/generate/form3674")
def generate_3674(data: ProjectMetadata):
    """Generate FDA Form 3674 (Certification of Compliance with ClinicalTrials.gov)"""
    try:
        # Generate the document using the template
        document_buffer = generate_form_3674(data.dict(exclude_none=True))
        
        # Return the document as a file download
        return StreamingResponse(
            document_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Form3674_{datetime.now().strftime('%Y%m%d')}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating Form 3674: {str(e)}")

# Generate Cover Letter
@app.post("/generate/cover-letter")
def generate_letter(data: ProjectMetadata):
    """Generate a cover letter for IND submission"""
    try:
        # Generate the document using the template
        document_buffer = generate_cover_letter(data.dict(exclude_none=True))
        
        # Return the document as a file download
        return StreamingResponse(
            document_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=CoverLetter_{datetime.now().strftime('%Y%m%d')}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating cover letter: {str(e)}")

# Generate Module 3 document
@app.post("/generate/module3")
def generate_module3(data: Module3Data):
    """Generate Module 3 (Chemistry, Manufacturing, and Controls) document"""
    try:
        # This is a stub for now - will be implemented later
        # Create a simple Word document as a placeholder
        from docx import Document
        doc = Document()
        doc.add_heading('Module 3: Chemistry, Manufacturing, and Controls', 0)
        doc.add_heading('Drug Substance Information', level=1)
        doc.add_paragraph(f'Drug Name: {data.drug_name}')
        doc.add_paragraph(f'Manufacturing Site: {data.manufacturing_site}')
        doc.add_paragraph(f'Batch Number: {data.batch_number}')
        
        doc.add_heading('Specifications', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Limit'
        hdr_cells[2].text = 'Result'
        
        for spec in data.specifications:
            row_cells = table.add_row().cells
            row_cells[0].text = spec.parameter
            row_cells[1].text = spec.limit
            row_cells[2].text = spec.result
        
        doc.add_heading('Stability Data', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Timepoint'
        hdr_cells[1].text = 'Result'
        
        for stability in data.stability_data:
            row_cells = table.add_row().cells
            row_cells[0].text = stability.timepoint
            row_cells[1].text = stability.result
        
        # Save the document to a BytesIO object
        document_buffer = BytesIO()
        doc.save(document_buffer)
        document_buffer.seek(0)
        
        # Return the document as a file download
        return StreamingResponse(
            document_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Module3_{datetime.now().strftime('%Y%m%d')}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating Module 3 document: {str(e)}")

# Direct run for debugging
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)