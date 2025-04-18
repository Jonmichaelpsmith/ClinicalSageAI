# create_template.py
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_module3_template():
    """Create a Module 3 CMC template document with Jinja2 placeholders"""
    
    # Create a new Document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Module 3 - Quality (CMC)"
    doc.core_properties.subject = "Investigational New Drug Application"
    doc.core_properties.author = "LumenTrialGuide.AI"
    
    # Add header with styling
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "CONFIDENTIAL - Investigational New Drug Application"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.runs[0]
    header_run.font.size = Pt(9)
    header_run.font.italic = True
    
    # Add title
    title = doc.add_heading("Module 3: Quality (Chemistry, Manufacturing and Controls)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add drug name section with template variable
    doc.add_heading("1. Drug Product Information", level=2)
    
    # Drug name paragraph with template variable
    name_para = doc.add_paragraph()
    name_para.add_run("Drug Name: ").bold = True
    name_para.add_run("{{ drug_name }}")
    
    # Nomenclature section
    doc.add_heading("1.1 Nomenclature", level=3)
    nom_para1 = doc.add_paragraph()
    nom_para1.add_run("Chemical Name: ").bold = True
    nom_para1.add_run("{{ nomenclature.chemical_name }}")
    
    nom_para2 = doc.add_paragraph()
    nom_para2.add_run("CAS Number: ").bold = True
    nom_para2.add_run("{{ nomenclature.cas_number }}")
    
    nom_para3 = doc.add_paragraph()
    nom_para3.add_run("Molecular Formula: ").bold = True
    nom_para3.add_run("{{ nomenclature.molecular_formula }}")
    
    # Manufacturing section
    doc.add_heading("2. Manufacturing Information", level=2)
    
    mfg_para1 = doc.add_paragraph()
    mfg_para1.add_run("Manufacturing Site: ").bold = True
    mfg_para1.add_run("{{ manufacturing_site }}")
    
    mfg_para2 = doc.add_paragraph()
    mfg_para2.add_run("Facility Address: ").bold = True
    mfg_para2.add_run("{{ facility_address }}")
    
    mfg_para3 = doc.add_paragraph()
    mfg_para3.add_run("Batch Number: ").bold = True
    mfg_para3.add_run("{{ batch_number }}")
    
    mfg_para4 = doc.add_paragraph()
    mfg_para4.add_run("Manufacture Date: ").bold = True
    mfg_para4.add_run("{{ manufacture_date }}")
    
    # Drug substance section
    doc.add_heading("3. Drug Substance Characterization", level=2)
    
    ds_para1 = doc.add_paragraph()
    ds_para1.add_run("Appearance: ").bold = True
    ds_para1.add_run("{{ drug_substance.appearance }}")
    
    ds_para2 = doc.add_paragraph()
    ds_para2.add_run("Solubility: ").bold = True
    ds_para2.add_run("{{ drug_substance.solubility }}")
    
    ds_para3 = doc.add_paragraph()
    ds_para3.add_run("Polymorphism: ").bold = True
    ds_para3.add_run("{{ drug_substance.polymorphism }}")
    
    ds_para4 = doc.add_paragraph()
    ds_para4.add_run("Synthesis Route: ").bold = True
    ds_para4.add_run("{{ drug_substance.synthesis_route }}")
    
    # Specifications section with table
    doc.add_heading("4. Specifications and Analytical Procedures", level=2)
    
    doc.add_paragraph("The following specifications have been established for the drug substance:")
    
    # Create specifications table
    spec_table = doc.add_table(rows=1, cols=4)
    spec_table.style = 'Table Grid'
    
    # Add header row
    header_cells = spec_table.rows[0].cells
    header_cells[0].text = "Parameter"
    header_cells[1].text = "Method"
    header_cells[2].text = "Acceptance Criteria"
    header_cells[3].text = "Result"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add Jinja2 template for-loop to add rows dynamically
    spec_table.add_row()
    spec_row = spec_table.rows[1].cells
    spec_row[0].text = "{% for spec in specifications %}\n{{ spec.parameter }}\n{% if not loop.last %}\n---\n{% endif %}\n{% endfor %}"
    spec_row[1].text = "{% for spec in specifications %}\n{{ spec.method }}\n{% if not loop.last %}\n---\n{% endif %}\n{% endfor %}"
    spec_row[2].text = "{% for spec in specifications %}\n{{ spec.acceptance_criteria }}\n{% if not loop.last %}\n---\n{% endif %}\n{% endfor %}"
    spec_row[3].text = "{% for spec in specifications %}\n{{ spec.result }}\n{% if not loop.last %}\n---\n{% endif %}\n{% endfor %}"
    
    # Stability data section
    doc.add_heading("5. Stability Data", level=2)
    
    doc.add_paragraph("The following stability data has been collected for the drug substance:")
    
    # Create stability table
    stability_table = doc.add_table(rows=1, cols=5)
    stability_table.style = 'Table Grid'
    
    # Add header row
    stability_header = stability_table.rows[0].cells
    stability_header[0].text = "Timepoint"
    stability_header[1].text = "Assay"
    stability_header[2].text = "Purity"
    stability_header[3].text = "Water Content"
    stability_header[4].text = "Appearance"
    
    # Make header row bold
    for cell in stability_header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add Jinja2 template for-loop to add rows dynamically
    stability_table.add_row()
    stability_row = stability_table.rows[1].cells
    stability_row[0].text = "{% for point in stability_data %}\n{{ point.timepoint }}\n{% if not loop.last %}\n---\n{% endif %}\n{% endfor %}"
    stability_row[1].text = "{% for point in stability_data %}\n{{ point.assay }}\n{% if not loop.last %}\n---\n{% endif %}\n{% endfor %}"
    stability_row[2].text = "{% for point in stability_data %}\n{{ point.purity }}\n{% if not loop.last %}\n---\n{% endif %}\n{% endfor %}"
    stability_row[3].text = "{% for point in stability_data %}\n{{ point.water_content }}\n{% if not loop.last %}\n---\n{% endif %}\n{% endfor %}"
    stability_row[4].text = "{% for point in stability_data %}\n{{ point.appearance }}\n{% if not loop.last %}\n---\n{% endif %}\n{% endfor %}"
    
    # Container closure
    doc.add_heading("6. Container Closure System", level=2)
    
    cc_para = doc.add_paragraph()
    cc_para.add_run("{{ container_closure }}")
    
    # Storage conditions
    doc.add_heading("7. Storage Conditions", level=2)
    
    sc_para = doc.add_paragraph()
    sc_para.add_run("{{ storage_conditions }}")
    
    # Save the document
    template_path = os.path.join('templates', 'module3_cmc.docx.j2')
    doc.save(template_path)
    
    print(f"Template created at {template_path}")

if __name__ == "__main__":
    # Ensure we're in the right directory
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    create_module3_template()