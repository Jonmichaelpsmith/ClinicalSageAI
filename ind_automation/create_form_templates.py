#!/usr/bin/env python3
"""
FDA Form Templates Generator

This script creates template DOCX files for FDA forms:
- Form 1571 (Investigational New Drug Application)
- Form 1572 (Statement of Investigator)
- Form 3674 (Certification of Compliance with ClinicalTrials.gov)
- Cover Letter for IND Submission

These templates will be used by the main.py FastAPI service to generate
customized forms based on provided data.
"""

import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# Create the templates directory if it doesn't exist
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), 'templates', 'forms')
os.makedirs(TEMPLATES_DIR, exist_ok=True)

def create_form_1571():
    """
    Create FDA Form 1571 (Investigational New Drug Application) template
    """
    doc = docx.Document()
    
    # Document properties
    doc.core_properties.title = "FDA Form 1571 - Investigational New Drug Application"
    doc.core_properties.keywords = "FDA, Form 1571, IND"
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("DEPARTMENT OF HEALTH AND HUMAN SERVICES\nFOOD AND DRUG ADMINISTRATION")
    header_run.bold = True
    header_run.font.size = Pt(12)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("INVESTIGATIONAL NEW DRUG APPLICATION (IND)\n(TITLE 21, CODE OF FEDERAL REGULATIONS (CFR) PART 312)")
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Form identification
    form_id_para = doc.add_paragraph()
    form_id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    form_id_run = form_id_para.add_run("Form FDA 1571 (07/18)")
    form_id_run.font.size = Pt(9)
    
    # Section 1
    section1_para = doc.add_paragraph()
    section1_para.style = 'List Number'
    section1_run = section1_para.add_run("NAME OF SPONSOR")
    section1_run.bold = True
    
    # Placeholder for sponsor name
    sponsor_para = doc.add_paragraph("{{sponsor_name}}")
    sponsor_para.paragraph_format.left_indent = Inches(0.5)
    
    # Section 2
    section2_para = doc.add_paragraph()
    section2_para.style = 'List Number'
    section2_run = section2_para.add_run("DATE OF SUBMISSION")
    section2_run.bold = True
    
    # Placeholder for submission date
    date_para = doc.add_paragraph("{{submission_date}}")
    date_para.paragraph_format.left_indent = Inches(0.5)
    
    # Section 3
    section3_para = doc.add_paragraph()
    section3_para.style = 'List Number'
    section3_run = section3_para.add_run("NAME OF DRUG (Include established and proprietary names)")
    section3_run.bold = True
    
    # Placeholder for drug name
    drug_para = doc.add_paragraph("{{drug_name}}")
    drug_para.paragraph_format.left_indent = Inches(0.5)
    
    # Section 4
    section4_para = doc.add_paragraph()
    section4_para.style = 'List Number'
    section4_run = section4_para.add_run("INDICATION(S) (Covered by this submission)")
    section4_run.bold = True
    
    # Placeholder for indication
    indication_para = doc.add_paragraph("{{indication}}")
    indication_para.paragraph_format.left_indent = Inches(0.5)
    
    # Section 5
    section5_para = doc.add_paragraph()
    section5_para.style = 'List Number'
    section5_run = section5_para.add_run("PHASE(S) OF CLINICAL INVESTIGATION TO BE CONDUCTED")
    section5_run.bold = True
    
    # Placeholder for phase
    phase_para = doc.add_paragraph("{{phase}}")
    phase_para.paragraph_format.left_indent = Inches(0.5)
    
    # Sections 6-11 with appropriate placeholders would be added here
    # ...
    
    # Section 12: Signature
    signature_section = doc.add_paragraph("\n\n")
    signature_section.add_run("NAME AND TITLE OF THE PERSON RESPONSIBLE FOR MONITORING THE CONDUCT AND PROGRESS OF THE CLINICAL INVESTIGATIONS")
    signature_section.add_run("\n{{contact_name}}, {{contact_title}}")
    
    signature_section.add_run("\n\nNAME, ADDRESS, AND TELEPHONE NUMBER OF PERSON TO WHOM QUESTIONS ABOUT THE APPLICATION SHOULD BE DIRECTED")
    signature_section.add_run("\n{{contact_name}}\n{{sponsor_address}}\n{{contact_phone}}\n{{contact_email}}")
    
    signature_section.add_run("\n\nSIGNATURE OF SPONSOR OR SPONSOR'S AUTHORIZED REPRESENTATIVE")
    signature_section.add_run("\n\n\n_____________________________________________    DATE: ____________________")
    
    signature_section.add_run("\n\nTYPED NAME AND TITLE")
    signature_section.add_run("\n\n{{authorizer_name}}, {{authorizer_title}}")
    
    # Save document
    doc.save(os.path.join(TEMPLATES_DIR, 'form1571.docx'))
    print(f"Created Form 1571 template at {os.path.join(TEMPLATES_DIR, 'form1571.docx')}")


def create_form_1572():
    """
    Create FDA Form 1572 (Statement of Investigator) template
    """
    doc = docx.Document()
    
    # Document properties
    doc.core_properties.title = "FDA Form 1572 - Statement of Investigator"
    doc.core_properties.keywords = "FDA, Form 1572, Investigator"
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("DEPARTMENT OF HEALTH AND HUMAN SERVICES\nFOOD AND DRUG ADMINISTRATION")
    header_run.bold = True
    header_run.font.size = Pt(12)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("STATEMENT OF INVESTIGATOR\n(TITLE 21, CODE OF FEDERAL REGULATIONS (CFR) PART 312)")
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Form identification
    form_id_para = doc.add_paragraph()
    form_id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    form_id_run = form_id_para.add_run("Form FDA 1572 (02/19)")
    form_id_run.font.size = Pt(9)
    
    # Introduction text
    intro_para = doc.add_paragraph()
    intro_text = "NOTE: No investigator may participate in an investigation until he/she provides the sponsor with a completed, signed Statement of Investigator, Form FDA 1572 (21 CFR 312.53(c))."
    intro_para.add_run(intro_text)
    
    # 1. Name and address of investigator
    section1 = doc.add_paragraph()
    section1.style = 'List Number'
    section1.add_run("NAME AND ADDRESS OF INVESTIGATOR").bold = True
    
    name_para = doc.add_paragraph("{{principal_investigator_name}}")
    name_para.paragraph_format.left_indent = Inches(0.5)
    
    address_para = doc.add_paragraph("{{investigator_address}}")
    address_para.paragraph_format.left_indent = Inches(0.5)
    
    # 2. Education, training, experience
    section2 = doc.add_paragraph()
    section2.style = 'List Number'
    section2.add_run("EDUCATION, TRAINING, AND EXPERIENCE THAT QUALIFIES THE INVESTIGATOR AS AN EXPERT IN THE CLINICAL INVESTIGATION OF THE DRUG FOR THE USE UNDER INVESTIGATION. ONE OF THE FOLLOWING IS ATTACHED:").bold = True
    
    cv_para = doc.add_paragraph("☐ CURRICULUM VITAE    ☐ OTHER STATEMENT OF QUALIFICATIONS")
    cv_para.paragraph_format.left_indent = Inches(0.5)
    
    # 3. Name and address of any medical school, etc.
    section3 = doc.add_paragraph()
    section3.style = 'List Number'
    section3.add_run("NAME AND ADDRESS OF ANY MEDICAL SCHOOL, HOSPITAL, OR OTHER RESEARCH FACILITY WHERE THE CLINICAL INVESTIGATION WILL BE CONDUCTED").bold = True
    
    facility_para = doc.add_paragraph("{{research_facility_name}}")
    facility_para.paragraph_format.left_indent = Inches(0.5)
    
    facility_address_para = doc.add_paragraph("{{research_facility_address}}")
    facility_address_para.paragraph_format.left_indent = Inches(0.5)
    
    # 4. Name and address of any clinical laboratory facilities
    section4 = doc.add_paragraph()
    section4.style = 'List Number'
    section4.add_run("NAME AND ADDRESS OF ANY CLINICAL LABORATORY FACILITIES TO BE USED IN THE STUDY").bold = True
    
    lab_para = doc.add_paragraph("{{clinical_lab_name}}")
    lab_para.paragraph_format.left_indent = Inches(0.5)
    
    lab_address_para = doc.add_paragraph("{{clinical_lab_address}}")
    lab_address_para.paragraph_format.left_indent = Inches(0.5)
    
    # 5. Name and address of IRB
    section5 = doc.add_paragraph()
    section5.style = 'List Number'
    section5.add_run("NAME AND ADDRESS OF INSTITUTIONAL REVIEW BOARD (IRB) RESPONSIBLE FOR REVIEW AND APPROVAL OF THE STUDY").bold = True
    
    irb_para = doc.add_paragraph("{{irb_name}}")
    irb_para.paragraph_format.left_indent = Inches(0.5)
    
    irb_address_para = doc.add_paragraph("{{irb_address}}")
    irb_address_para.paragraph_format.left_indent = Inches(0.5)
    
    # 6-8. Additional sections would be added here
    # ...
    
    # 9. Commitments section
    commitments = doc.add_paragraph("\nI agree to conduct the study in accordance with the relevant, current protocol and will only make changes in a protocol after notifying the sponsor, except when necessary to protect the safety, rights, or welfare of subjects.\n\n")
    commitments.add_run("I agree to personally conduct or supervise the described investigation.\n\n")
    commitments.add_run("I agree to inform any patients, or any persons used as controls, that the drugs are being used for investigational purposes...\n\n")
    
    # 10. Signature section
    signature_para = doc.add_paragraph("\n\n")
    signature_para.add_run("SIGNATURE OF INVESTIGATOR")
    signature_para.add_run("\n\n\n_____________________________________________    DATE: ____________________")
    
    # Save document
    doc.save(os.path.join(TEMPLATES_DIR, 'form1572.docx'))
    print(f"Created Form 1572 template at {os.path.join(TEMPLATES_DIR, 'form1572.docx')}")


def create_form_3674():
    """
    Create FDA Form 3674 (Certification of Compliance with ClinicalTrials.gov) template
    """
    doc = docx.Document()
    
    # Document properties
    doc.core_properties.title = "FDA Form 3674 - Certification of Compliance with ClinicalTrials.gov"
    doc.core_properties.keywords = "FDA, Form 3674, ClinicalTrials.gov"
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("DEPARTMENT OF HEALTH AND HUMAN SERVICES\nFOOD AND DRUG ADMINISTRATION")
    header_run.bold = True
    header_run.font.size = Pt(12)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("CERTIFICATION OF COMPLIANCE, UNDER 42 U.S.C. § 282(j)(5)(B), WITH REQUIREMENTS OF ClinicalTrials.gov DATA BANK")
    title_run.bold = True
    title_run.font.size = Pt(13)
    
    # Form identification
    form_id_para = doc.add_paragraph()
    form_id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    form_id_run = form_id_para.add_run("Form FDA 3674 (9/22)")
    form_id_run.font.size = Pt(9)
    
    # Introduction text
    intro_para = doc.add_paragraph()
    intro_text = "Public Law 110-85, enacted on September 27, 2007, amended the Public Health Service Act to require registration and reporting of data on certain clinical trials on the government Web site ClinicalTrials.gov."
    intro_para.add_run(intro_text)
    
    # Sponsor information
    sponsor_para = doc.add_paragraph("\nName of Sponsor or Applicant: {{sponsor_name}}")
    
    # Application information
    app_para = doc.add_paragraph()
    app_para.add_run("Application Type (check one):\n").bold = True
    app_para.add_run("☐ IND  ☐ NDA  ☐ BLA  ☐ PMA  ☐ HDE  ☐ 510(k)  ☐ de novo")
    
    app_number_para = doc.add_paragraph("\nApplication Number (if known): {{ind_number}}")
    product_para = doc.add_paragraph("\nProduct Name: {{drug_name}}")
    
    # Certification options
    cert_para = doc.add_paragraph()
    cert_para.add_run("\nCERTIFICATION (check one of the following boxes):").bold = True
    
    option1 = doc.add_paragraph("\n☐ I certify that the requirements of 42 U.S.C. § 282(j), section 402(j) of the Public Health Service Act, enacted by 121 Stat. 823, Public Law 110-85, do not apply because the application does not refer to a clinical trial.")
    
    option2 = doc.add_paragraph("\n☐ I certify that the requirements of 42 U.S.C. § 282(j), section 402(j) of the Public Health Service Act, enacted by 121 Stat. 823, Public Law 110-85, do not apply to any clinical trial referenced in the application.")
    
    option3 = doc.add_paragraph("\n☐ I certify that the requirements of 42 U.S.C. § 282(j), section 402(j) of the Public Health Service Act, enacted by 121 Stat. 823, Public Law 110-85, apply to one or more of the clinical trials referenced in the application and that those trials have been or will be registered as required by the statute.")
    
    option3_nct = doc.add_paragraph("\nNCT Numbers for Applicable Clinical Trials: {{nct_number}}")
    
    # Warning
    warning_para = doc.add_paragraph()
    warning_para.add_run("\nWARNING: A willfully and knowingly false statement is a criminal offense, U.S. Code, Title 18, Section 1001.").italic = True
    
    # Signature section
    signature_para = doc.add_paragraph("\n\n")
    signature_para.add_run("SIGNATURE OF SPONSOR/APPLICANT'S AUTHORIZED REPRESENTATIVE")
    signature_para.add_run("\n\n\n_____________________________________________    DATE: ____________________")
    
    signature_para.add_run("\n\nTyped Name and Title of the Certifying Official:")
    signature_para.add_run("\n\n{{certifier_name}}, {{certifier_title}}")
    
    # Save document
    doc.save(os.path.join(TEMPLATES_DIR, 'form3674.docx'))
    print(f"Created Form 3674 template at {os.path.join(TEMPLATES_DIR, 'form3674.docx')}")


def create_cover_letter():
    """
    Create cover letter template for IND submission
    """
    doc = docx.Document()
    
    # Document properties
    doc.core_properties.title = "Cover Letter for IND Submission"
    doc.core_properties.keywords = "IND, Cover Letter, FDA"
    
    # Date
    date_para = doc.add_paragraph("{{submission_date}}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # FDA Address
    fda_para = doc.add_paragraph("\nFood and Drug Administration\nCenter for Drug Evaluation and Research\nCentral Document Room\n5901-B Ammendale Road\nBeltsville, MD 20705-1266")
    
    # Subject line
    subject_para = doc.add_paragraph("\nSubject: ")
    subject_run = subject_para.add_run("Initial Investigational New Drug Application (IND)")
    subject_run.bold = True
    subject_para.add_run("\nDrug: {{drug_name}}")
    subject_para.add_run("\nIndication: {{indication}}")
    subject_para.add_run("\nSponsor: {{sponsor_name}}")
    subject_para.add_run("\nSerial Number: {{serial_number}}")
    
    # Greeting
    doc.add_paragraph("\nDear Sir or Madam:")
    
    # Body
    body_para = doc.add_paragraph("\n{{sponsor_name}} is pleased to submit this Initial Investigational New Drug Application for {{drug_name}} for the treatment of {{indication}}. ")
    body_para.add_run("This application is being submitted in accordance with 21 CFR Part 312.")
    
    # Content description
    doc.add_paragraph("\nThis submission contains the following:")
    
    modules_para = doc.add_paragraph()
    modules_para.paragraph_format.left_indent = Inches(0.5)
    modules_para.add_run("• FDA Form 1571\n")
    modules_para.add_run("• FDA Form 1572\n")
    modules_para.add_run("• FDA Form 3674\n")
    modules_para.add_run("• Investigator's Brochure\n")
    modules_para.add_run("• Clinical Protocol: {{protocol_number}} - {{protocol_title}}\n")
    modules_para.add_run("• Chemistry, Manufacturing, and Controls Information\n")
    modules_para.add_run("• Pharmacology and Toxicology Information\n")
    modules_para.add_run("• Previous Human Experience Information\n")
    
    # Contact information
    doc.add_paragraph("\nIf you have any questions or require additional information regarding this submission, please contact:")
    
    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.left_indent = Inches(0.5)
    contact_para.add_run("{{contact_name}}\n")
    contact_para.add_run("{{contact_title}}\n")
    contact_para.add_run("Phone: {{contact_phone}}\n")
    contact_para.add_run("Email: {{contact_email}}")
    
    # Closing
    doc.add_paragraph("\nSincerely,")
    
    doc.add_paragraph("\n\n\n____________________________")
    
    sig_para = doc.add_paragraph()
    sig_para.add_run("{{authorizer_name}}\n")
    sig_para.add_run("{{authorizer_title}}\n")
    sig_para.add_run("{{sponsor_name}}")
    
    # Save document
    doc.save(os.path.join(TEMPLATES_DIR, 'cover_letter.docx'))
    print(f"Created Cover Letter template at {os.path.join(TEMPLATES_DIR, 'cover_letter.docx')}")


def main():
    """
    Create all FDA form templates
    """
    print(f"Creating FDA form templates in {TEMPLATES_DIR}...")
    create_form_1571()
    create_form_1572()
    create_form_3674()
    create_cover_letter()
    print("All templates created successfully.")


if __name__ == "__main__":
    main()