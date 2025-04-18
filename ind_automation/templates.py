"""
IND Automation Templates Module

This module provides functions for generating FDA IND application forms and related documents
using python-docx instead of docxtpl.
"""

import os
import io
from typing import Dict, Any
from io import BytesIO
from datetime import datetime

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _get_template_path(template_name: str) -> str:
    """Get the absolute path to a template file"""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_dir, "templates", "forms", template_name)

def generate_form_1571(data: Dict[str, Any]) -> BytesIO:
    """
    Generate FDA Form 1571 (Investigational New Drug Application)
    
    Args:
        data: Dictionary containing form data
        
    Returns:
        BytesIO object containing the generated document
    """
    # Create a new document
    doc = docx.Document()
    
    # Add title
    title = doc.add_heading("DEPARTMENT OF HEALTH AND HUMAN SERVICES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading("FOOD AND DRUG ADMINISTRATION", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    form_title = doc.add_heading("INVESTIGATIONAL NEW DRUG APPLICATION (IND)", 1)
    form_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    form_number = doc.add_paragraph("FORM FDA 1571")
    form_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add form sections
    doc.add_heading("1. NAME OF SPONSOR", 2)
    doc.add_paragraph(data.get("sponsor_name", ""))
    
    doc.add_heading("2. DATE OF SUBMISSION", 2)
    doc.add_paragraph(data.get("submission_date", datetime.now().strftime("%Y-%m-%d")))
    
    doc.add_heading("3. ADDRESS (Number, Street, City, State, ZIP Code)", 2)
    doc.add_paragraph(data.get("sponsor_address", ""))
    
    doc.add_heading("4. TELEPHONE NUMBER", 2)
    doc.add_paragraph(data.get("sponsor_phone", ""))
    
    doc.add_heading("5. IND NUMBER (If previously assigned)", 2)
    doc.add_paragraph(data.get("ind_number", ""))
    
    doc.add_heading("6. INDICATION(S) (Covered by this submission)", 2)
    doc.add_paragraph(data.get("indication", ""))
    
    doc.add_heading("7. PHASE(S) OF CLINICAL INVESTIGATION TO BE CONDUCTED", 2)
    phase_para = doc.add_paragraph()
    phase = data.get("phase", "").upper()
    
    phase_para.add_run("☐ PHASE 1  ")
    if "1" in phase:
        phase_para.runs[-1].text = "☑ PHASE 1  "
    
    phase_para.add_run("☐ PHASE 2  ")
    if "2" in phase:
        phase_para.runs[-1].text = "☑ PHASE 2  "
    
    phase_para.add_run("☐ PHASE 3  ")
    if "3" in phase:
        phase_para.runs[-1].text = "☑ PHASE 3  "
    
    phase_para.add_run("☐ OTHER (Specify): _____________")
    if "OTHER" in phase:
        phase_para.runs[-1].text = "☑ OTHER (Specify): _____________"
    
    doc.add_heading("8. SUBMISSION CONTENTS", 2)
    
    submission_para = doc.add_paragraph()
    submission_para.add_run("☑ PROTOCOL  ")
    submission_para.add_run("☑ CHEMISTRY, MANUFACTURING, AND CONTROL DATA  ")
    submission_para.add_run("☑ PHARMACOLOGY AND TOXICOLOGY DATA  ")
    submission_para.add_run("☑ PREVIOUS HUMAN EXPERIENCE  ")
    
    doc.add_heading("9. SERIAL NUMBER ASSIGNED TO THIS SUBMISSION", 2)
    doc.add_paragraph(data.get("serial_number", "001"))
    
    doc.add_heading("10. THIS SUBMISSION CONTAINS THE FOLLOWING", 2)
    doc.add_paragraph("(Check all that apply)")
    
    contents_para = doc.add_paragraph()
    contents_para.add_run("☑ INITIAL INVESTIGATIONAL NEW DRUG APPLICATION (IND)  ")
    contents_para.add_run("☐ RESPONSE TO CLINICAL HOLD  ")
    contents_para.add_run("☐ ANNUAL REPORT  ")
    contents_para.add_run("☐ GENERAL CORRESPONDENCE  ")
    contents_para.add_run("☐ REQUEST FOR REINSTATEMENT OF IND THAT IS WITHDRAWN  ")
    contents_para.add_run("☐ DEVELOPMENT SAFETY UPDATE REPORT (DSUR)  ")
    
    # Add certification section
    doc.add_heading("11. CERTIFICATION", 2)
    certification_text = (
        "I agree not to begin clinical investigations until 30 days after FDA's receipt of the IND unless I receive "
        "earlier notification by FDA that the studies may begin. I also agree not to begin or continue clinical "
        "investigations covered by the IND if those studies are placed on clinical hold. I agree to conduct the "
        "investigation in accordance with all other applicable regulatory requirements."
    )
    doc.add_paragraph(certification_text)
    
    # Add signature block
    doc.add_paragraph("\n\n")
    signature_para = doc.add_paragraph("_______________________________")
    signature_para.add_run("\nSIGNATURE OF SPONSOR OR SPONSOR'S AUTHORIZED REPRESENTATIVE")
    
    name_para = doc.add_paragraph("\n\n")
    name_para.add_run(data.get("authorizer_name", ""))
    name_para.add_run("\nNAME AND TITLE OF THE PERSON SIGNING")
    
    # Save the document to a BytesIO object
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generate_form_1572(data: Dict[str, Any]) -> BytesIO:
    """
    Generate FDA Form 1572 (Statement of Investigator)
    
    Args:
        data: Dictionary containing form data
        
    Returns:
        BytesIO object containing the generated document
    """
    # Create a new document
    doc = docx.Document()
    
    # Add title
    title = doc.add_heading("DEPARTMENT OF HEALTH AND HUMAN SERVICES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading("FOOD AND DRUG ADMINISTRATION", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    form_title = doc.add_heading("STATEMENT OF INVESTIGATOR", 1)
    form_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    form_number = doc.add_paragraph("FORM FDA 1572")
    form_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add form sections
    doc.add_heading("1. NAME AND ADDRESS OF INVESTIGATOR", 2)
    doc.add_paragraph(data.get("principal_investigator_name", ""))
    doc.add_paragraph(data.get("investigator_address", ""))
    
    doc.add_heading("2. EDUCATION, TRAINING, AND EXPERIENCE THAT QUALIFIES THE INVESTIGATOR AS AN EXPERT", 2)
    doc.add_paragraph("☑ Curriculum Vitae attached")
    
    doc.add_heading("3. NAME AND ADDRESS OF ANY MEDICAL SCHOOL, HOSPITAL, OR OTHER RESEARCH FACILITY WHERE THE CLINICAL INVESTIGATION(S) WILL BE CONDUCTED", 2)
    doc.add_paragraph(data.get("research_facility_name", ""))
    doc.add_paragraph(data.get("research_facility_address", ""))
    
    doc.add_heading("4. NAME AND ADDRESS OF ANY CLINICAL LABORATORY FACILITIES TO BE USED IN THE STUDY", 2)
    doc.add_paragraph(data.get("clinical_lab_name", ""))
    doc.add_paragraph(data.get("clinical_lab_address", ""))
    
    doc.add_heading("5. NAME AND ADDRESS OF INSTITUTIONAL REVIEW BOARD (IRB) RESPONSIBLE FOR REVIEW AND APPROVAL OF THE STUDY(IES)", 2)
    doc.add_paragraph(data.get("irb_name", ""))
    doc.add_paragraph(data.get("irb_address", ""))
    
    doc.add_heading("6. NAMES OF THE SUBINVESTIGATORS WHO WILL BE ASSISTING THE INVESTIGATOR IN THE CONDUCT OF THE INVESTIGATION(S)", 2)
    doc.add_paragraph(data.get("subinvestigators", ""))
    
    doc.add_heading("7. NAME AND CODE NUMBER, IF ANY, OF THE PROTOCOL(S) IN THE IND FOR THE STUDY(IES) TO BE CONDUCTED BY THE INVESTIGATOR", 2)
    doc.add_paragraph(f"{data.get('protocol_title', '')}\nProtocol Number: {data.get('protocol_number', '')}")
    
    # Add commitments section
    doc.add_heading("8. COMMITMENTS", 2)
    commitments = [
        "I agree to conduct the study(ies) in accordance with the relevant, current protocol(s) and will only make changes in a protocol after notifying the sponsor, except when necessary to protect the safety, rights, or welfare of subjects.",
        "I agree to personally conduct or supervise the described investigation(s).",
        "I agree to inform any patients, or any persons used as controls, that the drugs are being used for investigational purposes and I will ensure that the requirements relating to obtaining informed consent in 21 CFR Part 50 and institutional review board (IRB) review and approval in 21 CFR Part 56 are met.",
        "I agree to report to the sponsor adverse experiences that occur in the course of the investigation(s) in accordance with 21 CFR 312.64.",
        "I have read and understand the information in the investigator's brochure, including the potential risks and side effects of the drug.",
        "I agree to ensure that all associates, colleagues, and employees assisting in the conduct of the study(ies) are informed about their obligations in meeting the above commitments.",
        "I agree to maintain adequate and accurate records in accordance with 21 CFR 312.62 and to make those records available for inspection in accordance with 21 CFR 312.68.",
        "I will ensure that an IRB that complies with the requirements of 21 CFR Part 56 will be responsible for the initial and continuing review and approval of the clinical investigation. I also agree to promptly report to the IRB all changes in the research activity and all unanticipated problems involving risks to human subjects or others. Additionally, I will not make any changes in the research without IRB approval, except where necessary to eliminate apparent immediate hazards to human subjects.",
        "I agree to comply with all other requirements regarding the obligations of clinical investigators and all other pertinent requirements in 21 CFR Part 312."
    ]
    
    for idx, commitment in enumerate(commitments, 1):
        doc.add_paragraph(f"{chr(96+idx)}. {commitment}")
    
    # Add signature block
    doc.add_paragraph("\n\n")
    signature_para = doc.add_paragraph("_______________________________")
    signature_para.add_run(f"\n{data.get('principal_investigator_name', '')}")
    signature_para.add_run("\nSIGNATURE OF INVESTIGATOR")
    
    date_para = doc.add_paragraph("\n\n")
    date_para.add_run(data.get("submission_date", datetime.now().strftime("%Y-%m-%d")))
    date_para.add_run("\nDATE")
    
    # Save the document to a BytesIO object
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generate_form_3674(data: Dict[str, Any]) -> BytesIO:
    """
    Generate FDA Form 3674 (Certification of Compliance with ClinicalTrials.gov)
    
    Args:
        data: Dictionary containing form data
        
    Returns:
        BytesIO object containing the generated document
    """
    # Create a new document
    doc = docx.Document()
    
    # Add title
    title = doc.add_heading("DEPARTMENT OF HEALTH AND HUMAN SERVICES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading("FOOD AND DRUG ADMINISTRATION", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    form_title = doc.add_heading("CERTIFICATION OF COMPLIANCE WITH REQUIREMENTS OF ClinicalTrials.gov DATA BANK", 1)
    form_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    form_number = doc.add_paragraph("FORM FDA 3674")
    form_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add form sections
    doc.add_heading("A. CERTIFICATION STATEMENT (Select one)", 2)
    
    certifications = [
        "The requirements of section 402(j) of the Public Health Service Act (PHS Act) as amended by Title VIII of the Food and Drug Administration Amendments Act of 2007 (FDAAA) apply to the investigation that is the subject of this submission.",
        "The requirements of section 402(j) of the PHS Act do not apply to any clinical trial referenced in this application or submission because the trial is not an 'applicable clinical trial' (see 42 U.S.C. § 282(j)(1)(A)(i)) under the statutory provisions.",
        "The requirements of section 402(j) of the PHS Act do not apply to any clinical trial referenced in this application or submission for one or more reasons.",
        "This application or submission does not reference any clinical trials that require submission of information under section 402(j) of the PHS Act."
    ]
    
    for idx, cert in enumerate(certifications, 1):
        certification_para = doc.add_paragraph()
        checked = "☐"
        if idx == 1:  # Default to the first option
            checked = "☑"
        certification_para.add_run(f"{checked} {idx}. {cert}")
    
    doc.add_heading("B. CLINICAL TRIAL INFORMATION", 2)
    doc.add_paragraph("NCT Number: " + data.get("nct_number", ""))
    doc.add_paragraph("Protocol Title: " + data.get("protocol_title", ""))
    doc.add_paragraph("Phase: " + data.get("phase", ""))
    
    # Add certification section
    doc.add_heading("C. CERTIFICATION", 2)
    doc.add_paragraph("I certify that the information provided above is true and correct.")
    
    # Add signature block
    doc.add_paragraph("\n\n")
    signature_para = doc.add_paragraph("_______________________________")
    signature_para.add_run("\nSIGNATURE OF CERTIFIER")
    
    name_title_para = doc.add_paragraph("\n\n")
    name_title_para.add_run(f"{data.get('certifier_name', '')}, {data.get('certifier_title', '')}")
    name_title_para.add_run("\nNAME AND TITLE OF THE PERSON SIGNING")
    
    contact_para = doc.add_paragraph("\n\n")
    contact_para.add_run(f"Address: {data.get('certifier_address', '')}")
    contact_para.add_run(f"\nEmail: {data.get('certifier_email', '')}")
    contact_para.add_run(f"\nPhone: {data.get('certifier_phone', '')}")
    
    date_para = doc.add_paragraph("\n\n")
    date_para.add_run(data.get("submission_date", datetime.now().strftime("%Y-%m-%d")))
    date_para.add_run("\nDATE")
    
    # Save the document to a BytesIO object
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generate_cover_letter(data: Dict[str, Any]) -> BytesIO:
    """
    Generate a cover letter for IND submission
    
    Args:
        data: Dictionary containing letter data
        
    Returns:
        BytesIO object containing the generated document
    """
    # Create a new document
    doc = docx.Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add sponsor info at the top
    doc.add_paragraph(data.get("sponsor_name", ""))
    doc.add_paragraph(data.get("sponsor_address", ""))
    doc.add_paragraph(data.get("submission_date", datetime.now().strftime("%B %d, %Y")))
    
    # Add FDA address
    doc.add_paragraph("\nFood and Drug Administration")
    doc.add_paragraph("Center for Drug Evaluation and Research")
    doc.add_paragraph("Central Document Room")
    doc.add_paragraph("5901-B Ammendale Road")
    doc.add_paragraph("Beltsville, MD 20705-1266")
    
    # Add subject line
    subject = doc.add_paragraph("\nSUBJECT: ")
    subject.add_run("INITIAL INVESTIGATIONAL NEW DRUG APPLICATION (IND)").bold = True
    subject.add_run(f"\nDRUG PRODUCT: {data.get('drug_name', '')}")
    subject.add_run(f"\nINDICATION: {data.get('indication', '')}")
    
    # Add salutation
    doc.add_paragraph("\nTo Whom It May Concern:")
    
    # Add body of letter
    doc.add_paragraph(f"Please find enclosed an Initial Investigational New Drug Application (IND) for {data.get('drug_name', '')}. This submission is being made by {data.get('sponsor_name', '')} to support the clinical investigation of {data.get('drug_name', '')} for the treatment of {data.get('indication', '')}.")
    
    doc.add_paragraph("The enclosed IND contains the following:")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run("Form FDA 1571 (Investigational New Drug Application)")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run("Table of Contents")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run("Introductory Statement and General Investigational Plan")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run("Investigator's Brochure")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run(f"Clinical Protocol: {data.get('protocol_title', '')}")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run("Chemistry, Manufacturing, and Controls Information")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run("Pharmacology and Toxicology Information")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run("Previous Human Experience")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run("Form FDA 1572 (Statement of Investigator)")
    
    contents = doc.add_paragraph()
    contents.style = 'List Bullet'
    contents.add_run("Form FDA 3674 (Certification of Compliance with ClinicalTrials.gov)")
    
    # Add closing paragraph
    doc.add_paragraph(f"\nWe look forward to working with the FDA on the development of {data.get('drug_name', '')}. If you have any questions or require additional information, please contact {data.get('contact_name', '')} at {data.get('contact_phone', '')} or via email at {data.get('contact_email', '')}.")
    
    # Add closing
    doc.add_paragraph("\nSincerely,")
    doc.add_paragraph("\n\n\n")
    doc.add_paragraph(f"{data.get('authorizer_name', '')}")
    doc.add_paragraph(f"{data.get('authorizer_title', '')}")
    doc.add_paragraph(f"{data.get('sponsor_name', '')}")
    
    # Save the document to a BytesIO object
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output